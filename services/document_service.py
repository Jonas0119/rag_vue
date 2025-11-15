"""
文档服务 - 文档上传、处理、管理
"""
import os
from typing import Optional, List, Tuple
from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document

from database import DocumentDAO
from utils.file_handler import (
    is_allowed_file, validate_file_size, generate_safe_filename,
    save_uploaded_file, delete_file, read_text_file, format_file_size
)
from utils.text_splitter import split_by_paragraphs
from utils.config import config
from .vector_store_service import get_vector_store_service


class DocumentService:
    """文档服务"""
    
    def __init__(self):
        self.doc_dao = DocumentDAO()
        self.vector_service = get_vector_store_service()
    
    def upload_document(self, user_id: int, uploaded_file) -> Tuple[bool, str]:
        """
        上传并处理文档
        
        Args:
            user_id: 用户 ID
            uploaded_file: Streamlit UploadedFile 对象
        
        Returns:
            (是否成功, 消息)
        """
        # 验证文件类型
        if not is_allowed_file(uploaded_file.name):
            return False, f"不支持的文件类型。支持的格式：{', '.join(['.pdf', '.txt', '.md', '.docx'])}"
        
        # 验证文件大小
        file_size = uploaded_file.size
        valid, error_msg = validate_file_size(file_size, config.MAX_FILE_SIZE)
        if not valid:
            return False, error_msg
        
        # 生成安全文件名
        safe_filename = generate_safe_filename(uploaded_file.name)
        user_dir = f"{config.USER_DATA_DIR}/user_{user_id}/uploads"
        filepath = os.path.join(user_dir, safe_filename)
        
        # 保存文件
        if not save_uploaded_file(uploaded_file, filepath):
            return False, "文件保存失败"
        
        # 获取文件扩展名
        file_ext = Path(uploaded_file.name).suffix.lower()
        
        # 创建文档记录
        try:
            doc_id = self.doc_dao.create_document(
                user_id=user_id,
                filename=safe_filename,
                original_filename=uploaded_file.name,
                filepath=filepath,
                file_size=file_size,
                file_type=file_ext,
                vector_collection=self.vector_service.get_collection_name(user_id)
            )
            
            # 异步处理文档（解析、分块、向量化）
            success, message = self._process_document(user_id, doc_id, filepath, file_ext)
            
            if success:
                return True, f"文档上传成功！共生成 {message} 个文本块"
            else:
                # 标记文档为错误状态
                self.doc_dao.mark_document_error(doc_id, message)
                return False, f"文档处理失败：{message}"
        
        except Exception as e:
            # 删除已保存的文件
            delete_file(filepath)
            return False, f"上传失败：{str(e)}"
    
    def _process_document(self, user_id: int, doc_id: str, filepath: str, file_type: str) -> Tuple[bool, str]:
        """
        处理文档：解析、分块、向量化
        
        Returns:
            (是否成功, 消息/块数量)
        """
        try:
            # 1. 解析文档
            if file_type == '.pdf':
                loader = PyPDFLoader(filepath)
                pages = loader.load()
                full_text = "\n\n".join([page.page_content for page in pages])
                page_count = len(pages)
            elif file_type in ['.txt', '.md']:
                full_text = read_text_file(filepath)
                if not full_text:
                    return False, "无法读取文件内容"
                page_count = None
            elif file_type == '.docx':
                # Word 文档处理
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(filepath)
                    full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                    page_count = None
                except Exception as e:
                    return False, f"Word 文档解析失败：{str(e)}"
            else:
                return False, f"不支持的文件类型：{file_type}"
            
            # 2. 段落级别分块
            chunks = split_by_paragraphs(full_text)
            
            if not chunks:
                return False, "文档内容为空或无法分块"
            
            # 3. 创建 Document 对象（带元数据）
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "doc_id": doc_id,
                        "chunk_id": i,
                        "user_id": user_id,
                        "source": filepath
                    }
                )
                documents.append(doc)
            
            # 4. 向量化并存入 Chroma
            self.vector_service.add_documents(user_id, documents)
            
            # 5. 更新文档状态
            self.doc_dao.mark_document_active(doc_id, len(chunks))
            
            # 更新页数（如果是 PDF）
            if page_count:
                self.doc_dao.update_document(doc_id, page_count=page_count)
            
            return True, str(len(chunks))
        
        except Exception as e:
            return False, str(e)
    
    def get_user_documents(self, user_id: int) -> List:
        """获取用户的所有文档"""
        docs = self.doc_dao.get_user_documents(user_id, status='active')
        
        # 转换为字典格式，添加格式化信息
        result = []
        for doc in docs:
            doc_dict = doc.to_dict()
            doc_dict['file_size_formatted'] = format_file_size(doc.file_size)
            result.append(doc_dict)
        
        return result
    
    def delete_document(self, user_id: int, doc_id: str) -> Tuple[bool, str]:
        """
        删除文档（包括文件和向量）
        
        Args:
            user_id: 用户 ID
            doc_id: 文档 ID
        
        Returns:
            (是否成功, 消息)
        """
        try:
            # 1. 获取文档信息
            doc = self.doc_dao.get_document(doc_id)
            if not doc:
                return False, "文档不存在"
            
            # 验证权限
            if doc.user_id != user_id:
                return False, "无权删除该文档"
            
            # 2. 从向量库删除
            self.vector_service.delete_documents(user_id, doc_id)
            
            # 3. 删除物理文件
            if os.path.exists(doc.filepath):
                delete_file(doc.filepath)
            
            # 4. 标记数据库记录为已删除
            self.doc_dao.delete_document(doc_id)
            
            return True, "文档已删除"
        
        except Exception as e:
            return False, f"删除失败：{str(e)}"
    
    def get_document_preview(self, user_id: int, doc_id: str, max_length: int = 1000) -> Optional[str]:
        """
        获取文档预览（前N个字符）
        
        Args:
            user_id: 用户 ID
            doc_id: 文档 ID
            max_length: 最大长度
        
        Returns:
            预览文本
        """
        try:
            doc = self.doc_dao.get_document(doc_id)
            if not doc or doc.user_id != user_id:
                return None
            
            # 根据文件类型读取内容
            content = None
            
            if doc.file_type == '.pdf':
                # PDF 文件使用 PyPDFLoader
                try:
                    loader = PyPDFLoader(doc.filepath)
                    pages = loader.load()
                    content = "\n\n".join([page.page_content for page in pages])
                except Exception as e:
                    return f"PDF 预览失败: {str(e)}"
            
            elif doc.file_type == '.docx':
                # Word 文档
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(doc.filepath)
                    content = "\n\n".join([para.text for para in docx_doc.paragraphs if para.text.strip()])
                except Exception as e:
                    return f"Word 文档预览失败: {str(e)}"
            
            elif doc.file_type in ['.txt', '.md']:
                # 文本文件
                content = read_text_file(doc.filepath)
            
            if not content:
                return "无法读取文档内容"
            
            # 截取前 N 个字符
            if len(content) > max_length:
                return content[:max_length] + "\n\n... (内容过长，仅显示前 {} 字符)".format(max_length)
            return content
        
        except Exception as e:
            return f"预览失败: {str(e)}"
    
    def get_user_stats(self, user_id: int) -> dict:
        """获取用户文档统计"""
        doc_count = self.doc_dao.get_document_count(user_id)
        storage_used = self.doc_dao.get_total_storage(user_id)
        vector_count = self.vector_service.get_document_count(user_id)
        
        return {
            'document_count': doc_count,
            'storage_used': storage_used,
            'storage_used_formatted': format_file_size(storage_used),
            'vector_count': vector_count
        }


# 全局文档服务实例
_document_service: Optional[DocumentService] = None


def get_document_service() -> DocumentService:
    """获取全局文档服务实例（单例）"""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service

